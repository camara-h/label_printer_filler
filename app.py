import copy
import io
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
try:
    import streamlit as st
except ModuleNotFoundError:
    st = None
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

APP_DIR = Path(__file__).parent
DEFAULT_TEMPLATE = APP_DIR / "Letter-125-NO0424.docx"
ROWS_PER_SHEET = 20
LABELS_PER_ROW_GROUP = 5
TOTAL_LABELS_PER_SHEET = ROWS_PER_SHEET * LABELS_PER_ROW_GROUP

ALIGNMENTS = {
    "Left": WD_ALIGN_PARAGRAPH.LEFT,
    "Center": WD_ALIGN_PARAGRAPH.CENTER,
    "Right": WD_ALIGN_PARAGRAPH.RIGHT,
}

DISPLAY_ALIGNMENTS = list(ALIGNMENTS.keys())

MIN_FONT_SIZE = 4.0
MAX_FONT_SIZE = 7.0
MIN_RECOMMENDED_FONT_SIZE = 5.0
MAX_CIRCLE_LINES = 3
MAX_RECTANGLE_LINES = 6
RECOMMENDED_RECTANGLE_LINES = 5


def label_to_table_columns(label_column: int) -> Tuple[int, int]:
    if not 1 <= int(label_column) <= LABELS_PER_ROW_GROUP:
        raise ValueError("Label column must be between 1 and 5.")
    circle_col = (int(label_column) - 1) * 3
    rectangle_col = circle_col + 1
    return circle_col, rectangle_col


def serialize_text(text: str, offset: int, enabled: bool) -> str:
    if not enabled:
        return text
    match = re.search(r"(\d+)(?!.*\d)", text or "")
    if not match:
        return text or ""
    number = match.group(1)
    value = int(number) + offset
    return (text or "")[: match.start()] + str(value).zfill(len(number)) + (text or "")[match.end() :]


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
    for tbl in list(cell.tables):
        t = tbl._element
        t.getparent().remove(t)


def set_cell_padding(cell, top="0", start="0", bottom="0", end="0"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_line_spacing(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "120")
    spacing.set(qn("w:lineRule"), "auto")


def add_tab_stop_right(paragraph, position_twips: int = 1200):
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(position_twips))
    tabs.append(tab)


def normalize_hex_color(value: Any, default: str = "#000000") -> str:
    text = str(value or default).strip()
    if not text.startswith("#"):
        text = "#" + text
    if re.fullmatch(r"#[0-9A-Fa-f]{6}", text):
        return text.upper()
    return default


def add_formatted_run(paragraph, text: str, font_size: float, bold: bool, color: str = "#000000"):
    run = paragraph.add_run(text or "")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(float(font_size))
    run.bold = bool(bold)
    color = normalize_hex_color(color)
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    return run


def write_cell_from_lines(cell, lines: List[Dict[str, Any]], label_offset: int = 0, override_left_texts: Optional[List[str]] = None, override_right_texts: Optional[List[str]] = None):
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_padding(cell, top="0", start="20", bottom="0", end="20")

    if not lines:
        cell.add_paragraph("")
        return

    for idx, line in enumerate(lines):
        paragraph = cell.add_paragraph()
        paragraph.alignment = ALIGNMENTS.get(line.get("align", "Center"), WD_ALIGN_PARAGRAPH.CENTER)
        set_line_spacing(paragraph)

        if override_left_texts is None:
            left_text = serialize_text(line.get("left_text", ""), label_offset, line.get("serialize_left", False))
        else:
            left_text = override_left_texts[idx] if idx < len(override_left_texts) else ""

        if override_right_texts is None:
            right_text = serialize_text(line.get("right_text", ""), label_offset, line.get("serialize_right", False))
        else:
            right_text = override_right_texts[idx] if idx < len(override_right_texts) else ""

        add_formatted_run(paragraph, left_text, line.get("font_size", 6.0), line.get("bold", False), line.get("color", "#000000"))
        if line.get("use_tab") and right_text:
            add_tab_stop_right(paragraph, int(line.get("tab_pos", 1200)))
            paragraph.add_run("\t")
            add_formatted_run(paragraph, right_text, line.get("font_size", 6.0), line.get("bold", False), line.get("color", "#000000"))


def cell_has_content(cell) -> bool:
    return bool(cell.text.strip())


def validate_template(doc: Document) -> List[str]:
    errors = []
    if len(doc.tables) < 1:
        errors.append("No table found in template.")
        return errors
    first_table = doc.tables[0]
    if len(first_table.rows) != ROWS_PER_SHEET:
        errors.append(f"Expected 20 rows in the first table, found {len(first_table.rows)}.")
    if len(first_table.columns) != 14:
        errors.append(f"Expected 14 columns in the first table, found {len(first_table.columns)}.")
    for i, table in enumerate(doc.tables, start=1):
        if len(table.rows) != ROWS_PER_SHEET or len(table.columns) != 14:
            errors.append(f"Table {i} does not match the expected 20 x 14 template structure.")
    return errors


def get_existing_occupied_positions(template_bytes: bytes) -> set:
    doc = Document(io.BytesIO(template_bytes))
    occupied = set()
    for sheet_idx, table in enumerate(doc.tables, start=1):
        if len(table.rows) != ROWS_PER_SHEET or len(table.columns) != 14:
            continue
        for row_idx in range(ROWS_PER_SHEET):
            for label_col in range(1, LABELS_PER_ROW_GROUP + 1):
                circle_col, rectangle_col = label_to_table_columns(label_col)
                if cell_has_content(table.cell(row_idx, circle_col)) or cell_has_content(table.cell(row_idx, rectangle_col)):
                    occupied.add((sheet_idx, row_idx + 1, label_col))
    return occupied


def make_blank_table_copy(table):
    new_tbl = copy.deepcopy(table._tbl)
    # Wrap the copied XML in a temporary document so python-docx can expose cells cleanly.
    # The copied XML is returned after clearing through a real table proxy in the target document.
    return new_tbl


def ensure_sheet_count(doc: Document, desired_sheets: int):
    if desired_sheets <= len(doc.tables):
        return
    if not doc.tables:
        raise ValueError("No table available to duplicate for additional pages.")

    source_table_xml = copy.deepcopy(doc.tables[0]._tbl)
    body = doc._body._element

    while len(doc.tables) < desired_sheets:
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run.append(br)
        paragraph.append(run)
        body.append(paragraph)
        body.append(copy.deepcopy(source_table_xml))

    # Newly duplicated pages must be blank, even if the source template was partially filled.
    for table in doc.tables[1:]:
        if len(table.rows) == ROWS_PER_SHEET and len(table.columns) == 14:
            for row_idx in range(ROWS_PER_SHEET):
                for col_idx in range(14):
                    clear_cell(table.cell(row_idx, col_idx))


def default_lines(kind: str) -> List[Dict[str, Any]]:
    if kind == "circle":
        return [
            {"left_text": "Tissue 1", "right_text": "", "use_tab": False, "font_size": 7.0, "bold": True, "align": "Center", "serialize_left": True, "serialize_right": False, "tab_pos": 700, "color": "#000000"},
            {"left_text": "EXP_ID", "right_text": "", "use_tab": False, "font_size": 5.0, "bold": False, "align": "Center", "serialize_left": False, "serialize_right": False, "tab_pos": 700, "color": "#000000"},
        ]
    return [
        {"left_text": "Tissue 1", "right_text": "", "use_tab": False, "font_size": 7.0, "bold": True, "align": "Center", "serialize_left": True, "serialize_right": False, "tab_pos": 1200, "color": "#000000"},
        {"left_text": "Tissue Biopsy", "right_text": "", "use_tab": False, "font_size": 6.0, "bold": False, "align": "Center", "serialize_left": False, "serialize_right": False, "tab_pos": 1200, "color": "#000000"},
        {"left_text": "EXP_ID", "right_text": "Exp_data", "use_tab": True, "font_size": 6.0, "bold": False, "align": "Right", "serialize_left": False, "serialize_right": False, "tab_pos": 1200, "color": "#000000"},
    ]


def new_label_set(name="Tissue", start_row=1, start_col=1, count=20) -> Dict[str, Any]:
    circle = default_lines("circle")
    rectangle = default_lines("rectangle")
    circle[0]["left_text"] = f"{name} 1"
    rectangle[0]["left_text"] = f"{name} 1"
    rectangle[1]["left_text"] = f"{name} Biopsy"
    return {
        "name": name,
        "start_sheet": 1,
        "start_row": start_row,
        "start_col": start_col,
        "count": count,
        "circle_lines": circle,
        "rectangle_lines": rectangle,
    }


def init_state():
    if "label_sets" not in st.session_state:
        st.session_state.label_sets = [new_label_set("Tissue", 1, 1, 20)]
    if "layout_df" not in st.session_state:
        st.session_state.layout_df = pd.DataFrame()


def normalize_lines(lines: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    clean = []
    for line in lines:
        clean.append({
            "left_text": str(line.get("left_text", "")),
            "right_text": str(line.get("right_text", "")),
            "use_tab": bool(line.get("use_tab", False)),
            "font_size": float(line.get("font_size", 6.0)),
            "bold": bool(line.get("bold", False)),
            "align": line.get("align", "Center") if line.get("align", "Center") in DISPLAY_ALIGNMENTS else "Center",
            "serialize_left": bool(line.get("serialize_left", False)),
            "serialize_right": bool(line.get("serialize_right", False)),
            "tab_pos": int(line.get("tab_pos", 1200)),
            "color": normalize_hex_color(line.get("color", "#000000")),
        })
    return clean


def sync_line_widget_state(prefix: str, lines: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Copy existing widget values back into the line data before add/remove actions.

    Streamlit reruns the full script when a button is clicked. If the button is
    placed before the line widgets, the app can otherwise rebuild from the old
    saved line data and temporarily hide the editor. This keeps the source list
    and visible widgets in sync on every rerun.
    """
    for idx, line in enumerate(lines):
        line["left_text"] = st.session_state.get(f"{prefix}_left_{idx}", line.get("left_text", ""))
        line["serialize_left"] = st.session_state.get(f"{prefix}_ser_left_{idx}", line.get("serialize_left", False))
        line["use_tab"] = st.session_state.get(f"{prefix}_tab_{idx}", line.get("use_tab", False))
        line["right_text"] = st.session_state.get(f"{prefix}_right_{idx}", line.get("right_text", ""))
        line["serialize_right"] = st.session_state.get(f"{prefix}_ser_right_{idx}", line.get("serialize_right", False))
        line["tab_pos"] = st.session_state.get(f"{prefix}_tabpos_{idx}", line.get("tab_pos", 1200))
        line["font_size"] = st.session_state.get(f"{prefix}_size_{idx}", line.get("font_size", 6.0))
        line["bold"] = st.session_state.get(f"{prefix}_bold_{idx}", line.get("bold", False))
        line["align"] = st.session_state.get(f"{prefix}_align_{idx}", line.get("align", "Center"))
        line["color"] = st.session_state.get(f"{prefix}_color_{idx}", line.get("color", "#000000"))
    return normalize_lines(lines)


def line_editor(prefix: str, label: str, lines: List[Dict[str, Any]], max_lines: Optional[int] = None, recommended_lines: Optional[int] = None) -> List[Dict[str, Any]]:
    lines = sync_line_widget_state(prefix, normalize_lines(lines))

    if max_lines is not None and len(lines) > max_lines:
        st.error(f"{label} labels can have a maximum of {max_lines} lines. Extra lines were removed because they are unlikely to print correctly.")
        lines = lines[:max_lines]

    if recommended_lines is not None and len(lines) > recommended_lines:
        st.warning(f"{label} labels usually print best with {recommended_lines} lines or fewer. You can use up to {max_lines} lines if needed.")

    c_add, c_remove = st.columns(2)
    with c_add:
        add_disabled = max_lines is not None and len(lines) >= max_lines
        if st.button(f"Add line to {label.lower()}", key=f"add_{prefix}", disabled=add_disabled):
            lines.append({"left_text": "", "right_text": "", "use_tab": False, "font_size": 6.0, "bold": False, "align": "Center", "serialize_left": False, "serialize_right": False, "tab_pos": 1000, "color": "#000000"})
            # Do not return early. Continue rendering so the new blank editor
            # appears immediately on the same rerun.
        if add_disabled:
            st.caption(f"Maximum reached: {label.lower()} labels are limited to {max_lines} lines.")
    with c_remove:
        if lines and st.button(f"Remove last {label.lower()} line", key=f"remove_{prefix}"):
            lines.pop()
            # Do not return early. Continue rendering the remaining editors.

    for idx, line in enumerate(lines):
        with st.expander(f"{label} line {idx + 1}", expanded=True):
            line["left_text"] = st.text_input("Text", value=line.get("left_text", ""), key=f"{prefix}_left_{idx}")
            line["serialize_left"] = st.checkbox("Serialize trailing number in this text", value=line.get("serialize_left", False), key=f"{prefix}_ser_left_{idx}")
            line["use_tab"] = st.checkbox("Add tab and right text on this same line", value=line.get("use_tab", False), key=f"{prefix}_tab_{idx}")
            if line["use_tab"]:
                line["right_text"] = st.text_input("Right text after tab", value=line.get("right_text", ""), key=f"{prefix}_right_{idx}")
                line["serialize_right"] = st.checkbox("Serialize trailing number in right text", value=line.get("serialize_right", False), key=f"{prefix}_ser_right_{idx}")
                line["tab_pos"] = st.number_input("Right tab position, twips", min_value=300, max_value=2200, value=int(line.get("tab_pos", 1200)), step=50, key=f"{prefix}_tabpos_{idx}")
            else:
                line["right_text"] = line.get("right_text", "")
                line["serialize_right"] = line.get("serialize_right", False)

            c1, c2, c3, c4 = st.columns(4)
            with c1:
                current_size = float(line.get("font_size", 6.0))
                current_size = min(MAX_FONT_SIZE, max(MIN_FONT_SIZE, current_size))
                line["font_size"] = st.number_input(
                    "Font size",
                    min_value=MIN_FONT_SIZE,
                    max_value=MAX_FONT_SIZE,
                    value=current_size,
                    step=0.5,
                    help="Recommended: 5 to 7 pt. 4 pt is allowed, but may be hard to read.",
                    key=f"{prefix}_size_{idx}",
                )
                if float(line["font_size"]) < MIN_RECOMMENDED_FONT_SIZE:
                    st.warning("4 pt may be difficult to read after printing.")
            with c2:
                line["bold"] = st.checkbox("Bold", value=bool(line.get("bold", False)), key=f"{prefix}_bold_{idx}")
            with c3:
                line["align"] = st.selectbox("Alignment", options=DISPLAY_ALIGNMENTS, index=DISPLAY_ALIGNMENTS.index(line.get("align", "Center")), key=f"{prefix}_align_{idx}")
            with c4:
                line["color"] = st.color_picker("Text color", value=normalize_hex_color(line.get("color", "#000000")), key=f"{prefix}_color_{idx}")
    return lines

def line_texts_for_label(lines: List[Dict[str, Any]], offset: int) -> Tuple[List[str], List[str], str]:
    lefts = []
    rights = []
    display = []
    for line in lines:
        left = serialize_text(line.get("left_text", ""), offset, line.get("serialize_left", False))
        right = serialize_text(line.get("right_text", ""), offset, line.get("serialize_right", False))
        lefts.append(left)
        rights.append(right)
        display.append(left + ((" | " + right) if line.get("use_tab") and right else ""))
    return lefts, rights, " / ".join(display)


def next_position(sheet: int, row: int, col: int) -> Tuple[int, int, int]:
    row += 1
    if row > ROWS_PER_SHEET:
        row = 1
        col += 1
    if col > LABELS_PER_ROW_GROUP:
        col = 1
        sheet += 1
    return sheet, row, col


def build_layout(label_sets: List[Dict[str, Any]], occupied: set, skip_occupied: bool) -> pd.DataFrame:
    rows = []
    planned = set()
    global_label_num = 1
    for set_idx, label_set in enumerate(label_sets):
        sheet = int(label_set.get("start_sheet", 1))
        row = int(label_set.get("start_row", 1))
        col = int(label_set.get("start_col", 1))
        count = int(label_set.get("count", 1))
        written_for_set = 0
        guard = 0
        while written_for_set < count:
            guard += 1
            if guard > count + 10000:
                raise ValueError("Could not find enough available labels. Please check blocked spaces and starting position.")
            candidate = (sheet, row, col)
            unavailable = candidate in planned or (skip_occupied and candidate in occupied)
            if not unavailable:
                circle_lefts, circle_rights, circle_display = line_texts_for_label(label_set["circle_lines"], written_for_set)
                rect_lefts, rect_rights, rect_display = line_texts_for_label(label_set["rectangle_lines"], written_for_set)
                rows.append({
                    "Use": True,
                    "Global #": global_label_num,
                    "Set #": set_idx + 1,
                    "Set name": label_set.get("name", f"Set {set_idx + 1}"),
                    "Within set #": written_for_set + 1,
                    "Sheet": sheet,
                    "Row": row,
                    "Label column": col,
                    "Circle text": circle_display,
                    "Rectangle text": rect_display,
                    "Circle left JSON": json.dumps(circle_lefts, ensure_ascii=False),
                    "Circle right JSON": json.dumps(circle_rights, ensure_ascii=False),
                    "Rectangle left JSON": json.dumps(rect_lefts, ensure_ascii=False),
                    "Rectangle right JSON": json.dumps(rect_rights, ensure_ascii=False),
                })
                planned.add(candidate)
                global_label_num += 1
                written_for_set += 1
            sheet, row, col = next_position(sheet, row, col)
    return pd.DataFrame(rows)


def layout_warnings(layout_df: pd.DataFrame, occupied: set) -> List[str]:
    warnings = []
    if layout_df.empty:
        return ["No labels were generated."]
    bad = layout_df[(layout_df["Sheet"] < 1) | (layout_df["Row"] < 1) | (layout_df["Row"] > 20) | (layout_df["Label column"] < 1) | (layout_df["Label column"] > 5)]
    if not bad.empty:
        warnings.append("Some edited positions are outside the valid range. Sheet must be at least 1, row must be 1 to 20, and label column must be 1 to 5.")
    pos_counts = layout_df[layout_df.get("Use", True)].groupby(["Sheet", "Row", "Label column"]).size().reset_index(name="n")
    duplicated = pos_counts[pos_counts["n"] > 1]
    if not duplicated.empty:
        preview = ", ".join([f"sheet {int(r.Sheet)}, row {int(r.Row)}, column {int(r['Label column'])}" for _, r in duplicated.head(10).iterrows()])
        warnings.append(f"Duplicate target positions in the editable layout: {preview}.")
    hits = []
    for _, row in layout_df[layout_df.get("Use", True)].iterrows():
        candidate = (int(row["Sheet"]), int(row["Row"]), int(row["Label column"]))
        if candidate in occupied:
            hits.append(candidate)
    if hits:
        preview = ", ".join([f"sheet {s}, row {r}, column {c}" for s, r, c in hits[:10]])
        warnings.append(f"Some target labels already contain text in the uploaded template: {preview}.")
    return warnings


def parse_json_list(value: Any) -> List[str]:
    if isinstance(value, list):
        return [str(x) for x in value]
    try:
        parsed = json.loads(value)
        if isinstance(parsed, list):
            return [str(x) for x in parsed]
    except Exception:
        pass
    return []


def fill_from_layout(template_bytes: bytes, label_sets: List[Dict[str, Any]], layout_df: pd.DataFrame, allow_overwrite: bool) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    errors = validate_template(doc)
    if errors:
        raise ValueError("Template validation failed: " + " ".join(errors))

    active_df = layout_df[layout_df.get("Use", True)].copy()
    if active_df.empty:
        raise ValueError("No active labels to write.")

    warnings = layout_warnings(active_df, get_existing_occupied_positions(template_bytes))
    blocking = [w for w in warnings if "outside the valid range" in w or "Duplicate target positions" in w]
    if blocking:
        raise ValueError(" ".join(blocking))
    if not allow_overwrite:
        occupied_warnings = [w for w in warnings if "already contain text" in w]
        if occupied_warnings:
            raise ValueError(occupied_warnings[0] + " Enable overwrite to continue.")

    max_sheet = int(active_df["Sheet"].max())
    ensure_sheet_count(doc, max_sheet)

    for _, layout_row in active_df.iterrows():
        set_idx = int(layout_row["Set #"]) - 1
        if set_idx < 0 or set_idx >= len(label_sets):
            raise ValueError("Editable layout refers to a label set that no longer exists. Rebuild the layout preview.")
        label_set = label_sets[set_idx]
        sheet = int(layout_row["Sheet"])
        row_num = int(layout_row["Row"])
        label_col = int(layout_row["Label column"])
        if sheet < 1 or row_num < 1 or row_num > 20 or label_col < 1 or label_col > 5:
            raise ValueError("All edited positions must be inside the valid sheet, row, and label column ranges.")

        table = doc.tables[sheet - 1]
        circle_col, rectangle_col = label_to_table_columns(label_col)

        circle_lefts = parse_json_list(layout_row.get("Circle left JSON", "[]"))
        circle_rights = parse_json_list(layout_row.get("Circle right JSON", "[]"))
        rect_lefts = parse_json_list(layout_row.get("Rectangle left JSON", "[]"))
        rect_rights = parse_json_list(layout_row.get("Rectangle right JSON", "[]"))

        write_cell_from_lines(table.cell(row_num - 1, circle_col), label_set["circle_lines"], override_left_texts=circle_lefts, override_right_texts=circle_rights)
        write_cell_from_lines(table.cell(row_num - 1, rectangle_col), label_set["rectangle_lines"], override_left_texts=rect_lefts, override_right_texts=rect_rights)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def layout_grid_html(layout_df: pd.DataFrame, occupied: set) -> str:
    active = layout_df[layout_df.get("Use", True)].copy() if not layout_df.empty else pd.DataFrame()
    if active.empty:
        return "<p>No layout generated yet.</p>"
    max_sheet = max(1, int(active["Sheet"].max()))
    html_parts = ["<style>.sheetgrid{border-collapse:collapse;margin-bottom:24px}.sheetgrid td,.sheetgrid th{border:1px solid #ddd;padding:4px;text-align:center;font-size:12px}.sheetgrid td{width:110px;height:34px}.used{background:#f7f7f7}.occupied{background:#ffe5e5}.planned{background:#e7f3ff}.conflict{background:#ffd3a8}.small{font-size:11px;color:#555}</style>"]
    pos_to_text = {}
    duplicates = set()
    for _, r in active.iterrows():
        key = (int(r["Sheet"]), int(r["Row"]), int(r["Label column"]))
        if key in pos_to_text:
            duplicates.add(key)
            pos_to_text[key] += f"<br>⚠ {r['Set name']} #{int(r['Within set #'])}"
        else:
            pos_to_text[key] = f"{r['Set name']} #{int(r['Within set #'])}"
    for sheet in range(1, max_sheet + 1):
        html_parts.append(f"<h4>Sheet {sheet}</h4><table class='sheetgrid'><tr><th>Row</th>" + "".join([f"<th>Label col {c}</th>" for c in range(1, 6)]) + "</tr>")
        for row in range(1, 21):
            html_parts.append(f"<tr><th>{row}</th>")
            for col in range(1, 6):
                key = (sheet, row, col)
                cls = "used"
                text = ""
                if key in occupied:
                    cls = "occupied"
                    text = "Existing text"
                if key in pos_to_text:
                    cls = "conflict" if key in duplicates or key in occupied else "planned"
                    text = pos_to_text[key]
                html_parts.append(f"<td class='{cls}'>{text}</td>")
            html_parts.append("</tr>")
        html_parts.append("</table>")
    return "".join(html_parts)


def main():
    st.set_page_config(page_title="LabTAG LCS-125WH Label Filler", layout="wide")
    init_state()

    st.title("LabTAG LCS-125WH Label Filler")
    st.caption("Uses the Word template as the source of truth and only writes formatted text into label cells.")

    with st.sidebar:
        st.header("Template")
        uploaded_template = st.file_uploader("Upload official or partially used .docx template", type=["docx"])
        use_default = st.checkbox("Use included LCS-125WH template", value=True)
        if uploaded_template is not None:
            template_bytes = uploaded_template.read()
            st.success("Using uploaded template.")
        elif use_default and DEFAULT_TEMPLATE.exists():
            template_bytes = DEFAULT_TEMPLATE.read_bytes()
            st.success("Using included template.")
        else:
            st.error("Upload a .docx template or keep the included template selected.")
            st.stop()

        try:
            template_doc = Document(io.BytesIO(template_bytes))
            template_errors = validate_template(template_doc)
            if template_errors:
                for error in template_errors:
                    st.error(error)
            existing_occupied = get_existing_occupied_positions(template_bytes)
            st.caption(f"Detected {len(existing_occupied)} occupied label positions in the current template.")
        except Exception as exc:
            st.error(f"Could not read template: {exc}")
            st.stop()

        st.header("Output behavior")
        skip_occupied = st.checkbox("Skip labels that already contain text", value=True)
        allow_overwrite = st.checkbox("Allow overwrite if edited layout targets used labels", value=False)
        st.caption("If more labels are requested than fit on the existing page, the app adds another blank copy of the template page.")

    st.subheader("1. Build label ID sets")
    top_cols = st.columns([1, 1, 2])
    with top_cols[0]:
        if st.button("Add another Label ID Set", type="secondary"):
            previous = copy.deepcopy(st.session_state.label_sets[-1])
            previous["name"] = f"Set {len(st.session_state.label_sets) + 1}"
            previous["start_col"] = min(5, int(previous.get("start_col", 1)) + 1)
            previous["start_row"] = 1
            st.session_state.label_sets.append(previous)
            st.rerun()
    with top_cols[1]:
        if len(st.session_state.label_sets) > 1 and st.button("Remove last set"):
            st.session_state.label_sets.pop()
            st.rerun()

    for set_idx, label_set in enumerate(st.session_state.label_sets):
        with st.expander(f"Label ID Set {set_idx + 1}: {label_set.get('name', '')}", expanded=True):
            c1, c2, c3, c4, c5 = st.columns(5)
            with c1:
                label_set["name"] = st.text_input("Set name", value=label_set.get("name", f"Set {set_idx + 1}"), key=f"set_name_{set_idx}")
            with c2:
                label_set["start_sheet"] = st.number_input("Start sheet", min_value=1, max_value=50, value=int(label_set.get("start_sheet", 1)), step=1, key=f"set_sheet_{set_idx}")
            with c3:
                label_set["start_row"] = st.number_input("Start row", min_value=1, max_value=20, value=int(label_set.get("start_row", 1)), step=1, key=f"set_row_{set_idx}")
            with c4:
                label_set["start_col"] = st.number_input("Start label column", min_value=1, max_value=5, value=int(label_set.get("start_col", 1)), step=1, key=f"set_col_{set_idx}")
            with c5:
                label_set["count"] = st.number_input("Labels to fill", min_value=1, max_value=1000, value=int(label_set.get("count", 20)), step=1, key=f"set_count_{set_idx}")

            ltab, rtab = st.tabs(["Circle formatting", "Rectangle formatting"])
            with ltab:
                label_set["circle_lines"] = line_editor(f"set{set_idx}_circle", "Circle", label_set.get("circle_lines", default_lines("circle")), max_lines=MAX_CIRCLE_LINES)
            with rtab:
                label_set["rectangle_lines"] = line_editor(f"set{set_idx}_rectangle", "Rectangle", label_set.get("rectangle_lines", default_lines("rectangle")), max_lines=MAX_RECTANGLE_LINES, recommended_lines=RECOMMENDED_RECTANGLE_LINES)

    st.divider()
    st.subheader("2. Build editable layout and preview")
    cbuild, cclear = st.columns([1, 3])
    with cbuild:
        if st.button("Build editable layout", type="primary"):
            try:
                st.session_state.layout_df = build_layout(st.session_state.label_sets, existing_occupied, skip_occupied)
                st.success("Editable layout generated.")
            except Exception as exc:
                st.error(str(exc))
    with cclear:
        st.caption("This creates the serialized rows first. Then you can manually move labels by editing Sheet, Row, and Label column, or fine tune the generated text JSON fields.")

    if not st.session_state.layout_df.empty:
        tab_preview, tab_grid, tab_advanced = st.tabs(["Editable layout", "Sheet map", "Advanced text editing"])
        with tab_preview:
            display_cols = ["Use", "Global #", "Set name", "Within set #", "Sheet", "Row", "Label column", "Circle text", "Rectangle text"]
            edited_display = st.data_editor(
                st.session_state.layout_df[display_cols],
                use_container_width=True,
                hide_index=True,
                num_rows="fixed",
                column_config={
                    "Use": st.column_config.CheckboxColumn("Use"),
                    "Sheet": st.column_config.NumberColumn("Sheet", min_value=1, step=1),
                    "Row": st.column_config.NumberColumn("Row", min_value=1, max_value=20, step=1),
                    "Label column": st.column_config.NumberColumn("Label column", min_value=1, max_value=5, step=1),
                },
                disabled=["Global #", "Set name", "Within set #", "Circle text", "Rectangle text"],
                key="layout_display_editor",
            )
            for col in ["Use", "Sheet", "Row", "Label column"]:
                st.session_state.layout_df[col] = edited_display[col]

            warnings = layout_warnings(st.session_state.layout_df, existing_occupied)
            if warnings:
                for warning in warnings:
                    if "already contain text" in warning and allow_overwrite:
                        st.warning(warning + " Overwrite is enabled.")
                    else:
                        st.warning(warning)
            else:
                st.success("No layout conflicts detected.")

        with tab_grid:
            st.markdown(layout_grid_html(st.session_state.layout_df, existing_occupied), unsafe_allow_html=True)
            st.caption("Blue means planned labels. Red means existing text from the uploaded template. Orange means a conflict or duplicate.")

        with tab_advanced:
            st.caption("The display preview is not enough to preserve line-level formatting. Edit these JSON lists only when you need to fine tune the final text after serialization.")
            hidden_cols = ["Global #", "Circle left JSON", "Circle right JSON", "Rectangle left JSON", "Rectangle right JSON"]
            edited_json = st.data_editor(
                st.session_state.layout_df[hidden_cols],
                use_container_width=True,
                hide_index=True,
                num_rows="fixed",
                disabled=["Global #"],
                key="json_layout_editor",
            )
            for col in hidden_cols[1:]:
                st.session_state.layout_df[col] = edited_json[col]

        st.divider()
        st.subheader("3. Generate DOCX")
        if st.button("Generate filled DOCX from edited layout", type="primary"):
            try:
                output_bytes = fill_from_layout(
                    template_bytes=template_bytes,
                    label_sets=copy.deepcopy(st.session_state.label_sets),
                    layout_df=st.session_state.layout_df,
                    allow_overwrite=allow_overwrite,
                )
                st.success("DOCX generated.")
                st.download_button(
                    label="Download filled template",
                    data=output_bytes,
                    file_name="filled_LCS_125WH_labels.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as exc:
                st.error(str(exc))
    else:
        st.info("Add one or more label sets, then click Build editable layout.")


if __name__ == "__main__":
    main()
